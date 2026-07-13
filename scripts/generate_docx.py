import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime

MD_FILE = Path(__file__).resolve().parents[1] / 'USER_MANUAL.md'
OUT_FILE = Path(__file__).resolve().parents[1] / 'USER_MANUAL_styled.docx'

if not MD_FILE.exists():
    print('Markdown source not found:', MD_FILE)
    sys.exit(1)

with MD_FILE.open('r', encoding='utf-8') as f:
    lines = [line.rstrip('\n') for line in f]

doc = Document()

# Cover page
sect = doc.sections[0]
sect.top_margin = Inches(1)

title = doc.add_paragraph()
title.alignment = 1  # center
run = title.add_run('NEXA POS System User Manual')
run.bold = True
run.font.size = Pt(28)

sub = doc.add_paragraph()
sub.alignment = 1
run = sub.add_run('User Guide and Operational Manual')
run.font.size = Pt(14)

date_p = doc.add_paragraph()
date_p.alignment = 1
run = date_p.add_run(datetime.now().strftime('%B %d, %Y'))
run.italic = True
run.font.size = Pt(10)

doc.add_page_break()

# Parse simple Markdown
i = 0
n = len(lines)

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)

while i < n:
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    # Headings
    if line.startswith('#'):
        # count hashes
        level = len(line) - len(line.lstrip('#'))
        text = line.lstrip('#').strip()
        if level == 1:
            doc.add_heading(text, level=1)
        elif level == 2:
            doc.add_heading(text, level=2)
        else:
            doc.add_heading(text, level=3)
        i += 1
        continue
    # Horizontal rule
    if line.startswith('---'):
        doc.add_page_break()
        i += 1
        continue
    # Bullet list
    if line.startswith('- '):
        # gather contiguous bullets
        while i < n and lines[i].strip().startswith('- '):
            item = lines[i].strip()[2:].strip()
            p = doc.add_paragraph(item, style='List Bullet')
            i += 1
        continue
    # Numbered list
    import re
    if re.match(r'^\d+\.\s+', line):
        while i < n and re.match(r'^\d+\.\s+', lines[i].strip()):
            item = re.sub(r'^\d+\.\s+', '', lines[i].strip())
            p = doc.add_paragraph(item, style='List Number')
            i += 1
        continue
    # Paragraph(s): collect until blank line
    para_lines = [lines[i]]
    i += 1
    while i < n and lines[i].strip() != '':
        para_lines.append(lines[i])
        i += 1
    text = '\n'.join(para_lines).strip()
    # Replace Markdown inline markers minimally
    text = text.replace('**', '').replace('*', '')
    add_paragraph(text)

# Save
try:
    doc.save(OUT_FILE)
    print('Saved:', OUT_FILE)
except Exception as e:
    print('Error saving:', e)
    sys.exit(2)
